"""
WordExporter —— 习题Word文档导出器。

功能：
  - 将习题导出为格式化的Word文档
  - 支持题目+答案分离排版
  - 支持单套/多套习题导出
  - 支持带批改结果的成绩报告导出

依赖：python-docx
"""

import os
from datetime import datetime
from typing import List, Optional
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


class WordExporter:
    """Word文档导出器。

    将习题、答案、成绩等数据导出为格式化的 .docx 文件。

    使用示例:
        >>> exporter = WordExporter()
        >>> exporter.export_exercise(exercise, "output.docx")
        >>> exporter.export_with_answers(exercise, answer_sheet, score, "report.docx")
    """

    def __init__(self, font_name: str = "宋体",
                 font_size: int = 14, title_size: int = 18,
                 columns: int = 5):
        """初始化导出器。

        Args:
            font_name: 正文字体名称。
            font_size: 正文字号。
            title_size: 标题字号。
            columns: 题目排列列数。
        """
        self._font_name = font_name
        self._font_size = font_size
        self._title_size = title_size
        self._columns = columns

    # ------------------------------------------------------------------
    # 公开API
    # ------------------------------------------------------------------

    def export_exercise(self, exercise, output_path: str,
                        include_answers: bool = False) -> str:
        """导出单套习题为Word文档。

        Args:
            exercise: Exercise 对象。
            output_path: 输出文件路径。
            include_answers: 是否包含答案页。

        Returns:
            输出文件的绝对路径。
        """
        doc = Document()
        self._setup_page(doc)

        # -- 标题 --
        self._add_title(doc, "口算练习题")

        # -- 信息栏 --
        self._add_info_line(doc, exercise)

        # -- 题目表格 --
        self._add_problem_table(doc, exercise.problems)

        # -- 答案页 --
        if include_answers:
            doc.add_page_break()
            self._add_title(doc, "参考答案")
            self._add_answer_table(doc, exercise.problems)

        # -- 页脚 --
        self._add_footer(doc)

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        return os.path.abspath(output_path)

    def export_with_answers(self, exercise, answer_sheet, score,
                            output_path: str) -> str:
        """导出带批改结果的习题报告。

        Args:
            exercise: Exercise 对象。
            answer_sheet: AnswerSheet 对象。
            score: Score 对象。
            output_path: 输出文件路径。

        Returns:
            输出文件的绝对路径。
        """
        doc = Document()
        self._setup_page(doc)

        # -- 标题 --
        self._add_title(doc, "口算练习成绩报告")

        # -- 学生信息 --
        self._add_student_info(doc, score)

        # -- 成绩摘要 --
        self._add_score_summary(doc, score)

        # -- 详细题目表格（含答案对比）--
        self._add_detailed_table(doc, exercise, answer_sheet, score)

        # -- 页脚 --
        self._add_footer(doc)

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        return os.path.abspath(output_path)

    def export_multi_exercise(self, exercises: list, output_path: str) -> str:
        """导出多套习题合集。

        Args:
            exercises: Exercise 对象列表。
            output_path: 输出文件路径。

        Returns:
            输出文件的绝对路径。
        """
        doc = Document()
        self._setup_page(doc)

        self._add_title(doc, "口算练习题合集")

        for i, ex in enumerate(exercises, 1):
            if i > 1:
                doc.add_page_break()
            self._add_subtitle(doc, f"第{i}套: {ex.exercise_id}")
            self._add_info_line(doc, ex)
            self._add_problem_table(doc, ex.problems)

        self._add_footer(doc)

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        return os.path.abspath(output_path)

    # ------------------------------------------------------------------
    # 内部方法
    # ------------------------------------------------------------------

    def _setup_page(self, doc: Document) -> None:
        """设置页面格式。"""
        section = doc.sections[0]
        section.page_width = Cm(21)   # A4
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = self._font_name
        font.size = Pt(self._font_size)
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self._font_name)

    def _add_title(self, doc: Document, text: str) -> None:
        """添加标题。"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.size = Pt(self._title_size)
        run.font.bold = True
        run.font.name = self._font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), self._font_name)

    def _add_subtitle(self, doc: Document, text: str) -> None:
        """添加副标题。"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = self._font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), self._font_name)

    def _add_info_line(self, doc: Document, exercise) -> None:
        """添加习题信息行（类型、题数、日期）。"""
        info_text = (
            f"类型: {exercise.exercise_type}    "
            f"题数: {exercise.problem_count} 题    "
            f"日期: {exercise.created_at.strftime('%Y-%m-%d')}"
        )
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(info_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph()  # 空行

    def _add_student_info(self, doc: Document, score) -> None:
        """添加学生信息。"""
        info = (
            f"学生: {score.student}    "
            f"日期: {score.graded_at.strftime('%Y-%m-%d %H:%M')}"
        )
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(info)
        run.font.size = Pt(11)
        doc.add_paragraph()

    def _add_score_summary(self, doc: Document, score) -> None:
        """添加成绩摘要。"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = (
            f"总题数: {score.total}    "
            f"正确: {score.correct}    "
            f"错误: {score.wrong}    "
            f"得分: {score.percentage}%"
        )
        run = para.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        color = (0, 128, 0) if score.percentage >= 80 else \
                (255, 165, 0) if score.percentage >= 60 else (255, 0, 0)
        run.font.color.rgb = RGBColor(*color)
        doc.add_paragraph()

    def _add_problem_table(self, doc: Document, problems: list) -> None:
        """添加题目表格（仅题目，无答案）。"""
        cols = self._columns
        rows = (len(problems) + cols - 1) // cols

        table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, p in enumerate(problems):
            row = i // cols
            col = i % cols
            cell = table.cell(row, col)
            # 题号 + 算式
            text = f"{i + 1:>3}. {p.left} {p.operator.symbol} {p.right} = ____"
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.font.size = Pt(self._font_size)
            run.font.name = self._font_name
            run.element.rPr.rFonts.set(qn('w:eastAsia'), self._font_name)

    def _add_answer_table(self, doc: Document, problems: list) -> None:
        """添加答案表格（题号 + 答案）。"""
        cols = self._columns
        rows = (len(problems) + cols - 1) // cols

        table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, p in enumerate(problems):
            row = i // cols
            col = i % cols
            cell = table.cell(row, col)
            text = f"{i + 1:>3}. {p.left} {p.operator.symbol} {p.right} = {p.answer}"
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.font.size = Pt(self._font_size)
            run.font.name = self._font_name
            run.element.rPr.rFonts.set(qn('w:eastAsia'), self._font_name)

    def _add_detailed_table(self, doc: Document, exercise, answer_sheet, score) -> None:
        """添加详细对比表（题目、学生答案、正确答案、对错）。"""
        table = doc.add_table(rows=len(exercise.problems) + 1, cols=5,
                              style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        headers = ["题号", "题目", "你的答案", "正确答案", "结果"]
        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(h)
            run.font.bold = True
            run.font.size = Pt(11)

        # 数据行
        wrong_set = set(score.wrong_indices)
        for i, p in enumerate(exercise.problems, 1):
            student_ans = answer_sheet.answers.get(i, "—")
            is_wrong = i in wrong_set

            data = [
                str(i),
                f"{p.left} {p.operator.symbol} {p.right}",
                str(student_ans),
                str(p.answer),
                "✗" if is_wrong else "✓",
            ]
            for j, val in enumerate(data):
                cell = table.cell(i, j)
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(val)
                run.font.size = Pt(11)
                if j == 4:  # 结果列
                    run.font.color.rgb = RGBColor(255, 0, 0) if is_wrong \
                        else RGBColor(0, 128, 0)

    def _add_footer(self, doc: Document) -> None:
        """添加页脚。"""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("口算练习系统 · 自动生成")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(150, 150, 150)
